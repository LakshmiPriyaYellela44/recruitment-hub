import io
import json
import logging
import asyncio
from typing import Dict, List, Optional
from PyPDF2 import PdfReader
from docx import Document
import google.genai
from app.aws_services.s3_client import S3Client
from app.core.config import settings

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parser for resume documents using Google Gemini LLM."""
    
    GEMINI_EXTRACTION_PROMPT = """You are an expert resume parser. Extract ALL relevant PROFESSIONAL SKILLS and information from the provided resume text and return a valid JSON object with the following structure.

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON - no markdown, no explanations, no comments
2. ALL skills must be categorized into ONE of: backend, cloud_devops, frontend, architecture, soft_skills, other
3. Extract EVERY PROFESSIONAL SKILL mentioned - be comprehensive
4. IMPORTANT: Extract ONLY professional skills and competencies, NOT development tools (IDEs, version control systems, or AI assistants)
5. DO NOT EXTRACT: VSCode, Visual Studio Code, Git, GitHub, Copilot, ChatGPT, JIRA, Confluence, Slack, or similar tools/IDEs
6. DO EXTRACT: Programming languages, frameworks, databases, cloud platforms, methodologies, and soft skills
7. Use proper capitalization (e.g., "Python" not "python", "AWS" not "aws")
8. Extract ALL experiences, educations, certifications mentioned
9. For dates: extract whatever format is provided
10. DO NOT make up or assume data - only extract what's explicitly in the resume
11. Return empty arrays [] if no data found, never return null for arrays

Return EXACTLY this JSON structure:
{{
  "name": "Full name",
  "email": "Email or null",
  "phone": "Phone or null",
  "summary": "Professional summary (max 300 chars) or null",
  "skills": {{
    "backend": ["Python", "Django", "FastAPI", "PostgreSQL"],
    "cloud_devops": ["AWS", "Docker", "Kubernetes"],
    "frontend": ["React", "JavaScript", "TypeScript"],
    "architecture": ["System Design", "Microservices"],
    "soft_skills": ["Communication", "Problem-solving"],
    "other": []
  }},
  "experiences": [
    {{
      "title": "Job title",
      "company": "Company name",
      "location": "City, Country or null",
      "start_date": "Month Year or null",
      "end_date": "Month Year or 'Present' or null",
      "duration_description": "e.g., '3 years 6 months' or null",
      "description": "Summary of role and achievements or null"
    }}
  ],
  "educations": [
    {{
      "degree": "Full degree name",
      "field": "Field of study",
      "institution": "University name",
      "graduation_year": 2020,
      "gpa": null,
      "additional_info": "Honors, awards, etc. or null"
    }}
  ],
  "certifications": [
    {{
      "name": "Certification name",
      "issuer": "Issuing organization",
      "date": "Date obtained or null"
    }}
  ],
  "languages": ["language1", "language2"],
  "projects": [
    {{
      "name": "Project name",
      "description": "What was built",
      "technologies": ["tech1", "tech2"]
    }}
  ],
  "additional_info": "Other relevant information or null"
}}

SKILL CATEGORIZATION RULES:
- Backend: Python, Java, C#, Node.js, Django, FastAPI, Spring Boot, PostgreSQL, MongoDB, MySQL, Redis, etc.
- Cloud/DevOps: AWS, Azure, Google Cloud, Docker, Kubernetes, Jenkins, Terraform, CI/CD, etc.
- Frontend: React, Vue, Angular, JavaScript, TypeScript, HTML, CSS, Tailwind, Next.js, etc.
- Architecture: System Design, Microservices, REST APIs, GraphQL, Event-driven, SOLID, Design Patterns, etc.
- Soft Skills: Communication, Leadership, Teamwork, Problem-solving, Adaptability, Critical thinking, etc.
- Other: Any other legitimate professional tools or frameworks

TOOLS TO IGNORE (DO NOT EXTRACT):
- IDEs: VSCode, Visual Studio Code, IntelliJ, Sublime Text, Eclipse
- Version Control: Git, GitHub, GitLab, Bitbucket, SVN
- AI Assistants: Copilot, ChatGPT, GPT, Claude
- Collaboration: JIRA, Confluence, Slack
- Browser: Chrome, Firefox, Safari, Edge

Resume Text:
{resume_text}"""

    # List of tools/IDEs/utilities to EXCLUDE from skills (not professional competencies)
    NON_SKILLS = {
        # IDEs
        'vscode', 'visual studio code', 'vs code', 'intelij', 'intellij idea', 'sublime', 'sublime text',
        'eclipse', 'netbeans', 'pycharm', 'webstorm', 'xcode', 'android studio',
        # Version Control
        'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial', 'perforce',
        # AI Assistants
        'copilot', 'chatgpt', 'gpt', 'openai', 'claude', 'ai assistant', 'ai tools',
        # Collaboration/Project Management
        'jira', 'confluence', 'slack', 'asana', 'monday.com', 'trello', 'notion',
        # Office Tools
        'word', 'excel', 'powerpoint', 'outlook', 'microsoft office', 'google docs', 'sheets',
        # Browsers
        'chrome', 'firefox', 'safari', 'edge', 'internet explorer', 'browser',
        # Generic utilities
        'terminal', 'cmd', 'command line', 'shell', 'powershell', 'cmd.exe',
        # Other non-skills
        'linux', 'windows', 'macos', 'mac os', 'ubuntu', 'centos', 'os',
        'npm', 'yarn', 'pip', 'maven', 'gradle', 'webpack',  # Package managers/build tools (keep frameworks but not build tools)
        'postman', 'insomnia', 'figma', 'sketch', 'adobe xd',  # Design/API testing tools
        'zoom', 'teams', 'google meet', 'discord',  # Communication tools
    }

    @staticmethod
    def _filter_non_skills(skills: Dict[str, List[str]]) -> Dict[str, List[str]]:
        """
        Remove development tools/IDEs from skills while preserving ALL genuine professional skills.
        
        Args:
            skills: Dictionary with skill categories as keys and lists of skills as values
            
        Returns:
            Filtered skills dictionary with non-skills removed
        """
        filtered = {}
        
        for category, skill_list in skills.items():
            if not isinstance(skill_list, list):
                filtered[category] = skill_list
                continue
            
            # Filter out non-skills by checking against NON_SKILLS set
            # Normalize: lowercase, strip whitespace, replace hyphens/underscores with spaces
            filtered_list = []
            removed = []
            
            for skill in skill_list:
                if not skill or skill.strip() == '':
                    continue
                
                # Normalize skill name for comparison
                normalized_skill = skill.lower().strip()
                normalized_skill = normalized_skill.replace('_', ' ').replace('-', ' ')
                # Collapse multiple spaces to single space
                normalized_skill = ' '.join(normalized_skill.split())
                
                # Check if this normalized skill is in the non-skills list
                is_non_skill = normalized_skill in ResumeParser.NON_SKILLS
                
                if is_non_skill:
                    removed.append(skill)
                else:
                    filtered_list.append(skill)
            
            filtered[category] = filtered_list
            
            # Log removed skills for transparency
            if removed:
                logger.info(f"[_filter_non_skills] Removed {len(removed)} non-skills from '{category}': {removed}")
            if removed:
                logger.info(f"[_filter_non_skills] Removed non-skills from '{category}': {removed}")
        
        return filtered
    
    @staticmethod
    def _is_s3_key(file_path: str) -> bool:
        """Detect if file_path is an S3 key (not a local path)."""
        if '\\' in file_path:  # Windows path
            return False
        if file_path.startswith('/'):  # Unix path
            return False
        if len(file_path) > 1 and file_path[1] == ':':  # Drive letter
            return False
        return True
    
    @staticmethod
    async def _download_from_s3(s3_key: str) -> Optional[bytes]:
        """Download file from S3 with retries."""
        max_retries = 3
        retry_delay = 1
        
        for attempt in range(max_retries):
            try:
                s3_client = S3Client()
                logger.info(f"[_download_from_s3] Downloading from S3 (attempt {attempt + 1}/{max_retries}): {s3_key}")
                content = await s3_client.download_file(s3_key)
                if content is None:
                    logger.warning(f"[_download_from_s3] File not found in S3: {s3_key}")
                    return None
                logger.info(f"[_download_from_s3] Downloaded {len(content)} bytes successfully")
                return content
            except Exception as e:
                logger.warning(f"[_download_from_s3] Attempt {attempt + 1}/{max_retries} failed: {str(e)}")
                if attempt < max_retries - 1:
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2
                else:
                    logger.error(f"[_download_from_s3] Max retries exceeded: {str(e)}")
                    raise
        
        return None
    
    @staticmethod
    async def parse_pdf(file_path: str) -> Dict:
        """Parse PDF resume using Gemini LLM."""
        try:
            logger.info(f"[parse_pdf] Starting PDF parse for file_path: {file_path[:100]}")
            
            is_s3 = ResumeParser._is_s3_key(file_path)
            logger.info(f"[parse_pdf] S3 key detection: is_s3={is_s3}")
            
            if is_s3:
                logger.info(f"[parse_pdf] Downloading from S3")
                pdf_content = await ResumeParser._download_from_s3(file_path)
                if pdf_content is None or len(pdf_content) == 0:
                    error_msg = f"Could not download or empty PDF from S3: {file_path}"
                    logger.error(f"[parse_pdf] {error_msg}")
                    return {"error": error_msg}
                file_obj = io.BytesIO(pdf_content)
            else:
                logger.info(f"[parse_pdf] Reading local file: {file_path}")
                file_obj = open(file_path, "rb")
            
            # Extract text from PDF
            text = ""
            try:
                reader = PdfReader(file_obj)
                if len(reader.pages) == 0:
                    error_msg = f"PDF file has no pages: {file_path}"
                    logger.error(f"[parse_pdf] {error_msg}")
                    if not isinstance(file_obj, io.BytesIO):
                        file_obj.close()
                    return {"error": error_msg}
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                    except Exception as page_err:
                        logger.warning(f"[parse_pdf] Failed to extract page {page_num}: {str(page_err)}")
                        continue
            except Exception as pdf_err:
                error_msg = f"Failed to read PDF: {str(pdf_err)}"
                logger.error(f"[parse_pdf] {error_msg}")
                if not isinstance(file_obj, io.BytesIO):
                    file_obj.close()
                return {"error": error_msg}
            finally:
                if not isinstance(file_obj, io.BytesIO):
                    file_obj.close()
            
            if not text or len(text.strip()) == 0:
                error_msg = f"No text extracted from PDF: {file_path}"
                logger.error(f"[parse_pdf] {error_msg}")
                return {"error": error_msg}
            
            logger.info(f"[parse_pdf] Extracted {len(text)} characters from PDF")
            return await ResumeParser._parse_with_gemini(text)
        except Exception as e:
            logger.error(f"[parse_pdf] Exception: {str(e)}", exc_info=True)
            return {"error": f"PDF parsing failed: {str(e)}"}
    
    @staticmethod
    async def parse_docx(file_path: str) -> Dict:
        """Parse DOCX resume using Gemini LLM."""
        try:
            logger.info(f"[parse_docx] Starting DOCX parse for file_path: {file_path[:100]}")
            
            is_s3 = ResumeParser._is_s3_key(file_path)
            logger.info(f"[parse_docx] S3 key detection: is_s3={is_s3}")
            
            if is_s3:
                logger.info(f"[parse_docx] Downloading from S3")
                docx_content = await ResumeParser._download_from_s3(file_path)
                if docx_content is None or len(docx_content) == 0:
                    error_msg = f"Could not download or empty DOCX from S3: {file_path}"
                    logger.error(f"[parse_docx] {error_msg}")
                    return {"error": error_msg}
                file_obj = io.BytesIO(docx_content)
            else:
                logger.info(f"[parse_docx] Reading local file: {file_path}")
                file_obj = open(file_path, "rb")
            
            # Extract text from DOCX
            try:
                doc = Document(file_obj)
                if not doc.paragraphs:
                    error_msg = f"DOCX file has no paragraphs: {file_path}"
                    logger.error(f"[parse_docx] {error_msg}")
                    if not isinstance(file_obj, io.BytesIO):
                        file_obj.close()
                    return {"error": error_msg}
                
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as docx_err:
                error_msg = f"Failed to read DOCX: {str(docx_err)}"
                logger.error(f"[parse_docx] {error_msg}")
                if not isinstance(file_obj, io.BytesIO):
                    file_obj.close()
                return {"error": error_msg}
            finally:
                if not isinstance(file_obj, io.BytesIO):
                    file_obj.close()
            
            if not text or len(text.strip()) == 0:
                error_msg = f"No text extracted from DOCX: {file_path}"
                logger.error(f"[parse_docx] {error_msg}")
                return {"error": error_msg}
            
            logger.info(f"[parse_docx] Extracted {len(text)} characters from DOCX")
            return await ResumeParser._parse_with_gemini(text)
        except Exception as e:
            logger.error(f"[parse_docx] Exception: {str(e)}", exc_info=True)
            return {"error": f"DOCX parsing failed: {str(e)}"}
    
    @staticmethod
    async def _parse_with_gemini(resume_text: str) -> Dict:
        """Use Gemini LLM to extract resume data."""
        try:
            if not settings.GEMINI_API_KEY:
                error_msg = "GEMINI_API_KEY not configured. LLM extraction is required."
                logger.error(f"[_parse_with_gemini] {error_msg}")
                return {"error": error_msg}
            
            logger.info(f"[_parse_with_gemini] Starting Gemini extraction for {len(resume_text)} chars")
            
            # Prepare prompt
            prompt = ResumeParser.GEMINI_EXTRACTION_PROMPT.format(resume_text=resume_text)
            
            # Call Gemini API
            logger.info(f"[_parse_with_gemini] Calling Gemini API with model: {settings.GEMINI_MODEL}")
            client = google.genai.Client(api_key=settings.GEMINI_API_KEY)
            response = client.models.generate_content(
                model=settings.GEMINI_MODEL,
                contents=prompt,
                config=google.genai.types.GenerateContentConfig(
                    temperature=0.1,  # Low temperature for deterministic extraction
                    top_p=0.9,
                    max_output_tokens=4096
                )
            )
            
            if not response or not response.text:
                error_msg = "Gemini returned empty response"
                logger.error(f"[_parse_with_gemini] {error_msg}")
                return {"error": error_msg}
            
            logger.info(f"[_parse_with_gemini] Gemini response received: {len(response.text)} chars")
            logger.debug(f"[_parse_with_gemini] Raw response: {response.text[:500]}")
            
            # Parse JSON response
            extracted_data = ResumeParser._parse_json_response(response.text)
            
            # Filter out non-skills (tools, IDEs, etc.) while keeping professional competencies
            if 'skills' in extracted_data and isinstance(extracted_data['skills'], dict):
                extracted_data['skills'] = ResumeParser._filter_non_skills(extracted_data['skills'])
            
            logger.info(f"[_parse_with_gemini] Extraction complete:")
            logger.info(f"  - Name: {extracted_data.get('name')}")
            logger.info(f"  - Email: {extracted_data.get('email')}")
            
            # Log skill categories and counts
            if isinstance(extracted_data.get('skills'), dict):
                skills_summary = {cat: len(skills) for cat, skills in extracted_data['skills'].items() if skills}
                logger.info(f"  - Skills by category: {skills_summary}")
            else:
                logger.info(f"  - Skills: {len(extracted_data.get('skills', []))} found")
            
            logger.info(f"  - Experiences: {len(extracted_data.get('experiences', []))} found")
            logger.info(f"  - Educations: {len(extracted_data.get('educations', []))} found")
            logger.info(f"  - Certifications: {len(extracted_data.get('certifications', []))} found")
            
            return extracted_data
        except Exception as e:
            logger.error(f"[_parse_with_gemini] Exception: {str(e)}", exc_info=True)
            return {"error": f"Gemini parsing failed: {str(e)}"}
    

    @staticmethod
    def _parse_json_response(response_text: str) -> Dict:
        """Parse JSON from Gemini response, handling markdown formatting."""
        try:
            # Try direct JSON parsing first
            try:
                return json.loads(response_text)
            except json.JSONDecodeError:
                pass
            
            # Try extracting JSON from markdown code blocks
            if "```json" in response_text:
                json_str = response_text.split("```json")[1].split("```")[0].strip()
                return json.loads(json_str)
            
            if "```" in response_text:
                json_str = response_text.split("```")[1].split("```")[0].strip()
                return json.loads(json_str)
            
            # If still no JSON, try to find JSON object
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
            
            logger.error(f"[_parse_json_response] Could not extract JSON from: {response_text[:200]}")
            return {"error": "Failed to parse Gemini response as JSON"}
        except Exception as e:
            logger.error(f"[_parse_json_response] Exception: {str(e)}")
            return {"error": f"JSON parsing failed: {str(e)}"}

