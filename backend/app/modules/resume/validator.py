"""
Resume Validation Module - Comprehensive checks for resume authenticity and quality
"""
import io
import logging
import re
from typing import Dict, Tuple, List
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class ResumeValidator:
    """Validates resume files for authenticity, quality, and genuineness."""

    # Minimum required quality thresholds
    MIN_EXTRACTED_FIELDS = 2  # Must extract at least 2 meaningful fields
    MIN_TEXT_LENGTH = 200  # Minimum characters for resume content
    MAX_TEXT_LENGTH = 500000  # Maximum reasonable resume size
    MIN_SKILLS_OR_EXPERIENCES = 1  # Must have at least skills OR experiences
    
    # Common placeholder/template keywords that indicate fake resume
    FAKE_RESUME_KEYWORDS = [
        "your name here",
        "your address",
        "your email",
        "your phone",
        "[your",
        "{your",
        "sample resume",
        "template",
        "lorem ipsum",
        "john doe",
        "jane smith",
        "xyz company",
        "dummy",
        "example@example",
        "xxx-xxxx",
        "placeholder"
    ]
    
    # Suspicious patterns
    SUSPICIOUS_PATTERNS = {
        "too_many_dates": r"\d{4}[\s\-]\d{4}",  # Many year ranges
        "repeated_text": r"(.{10,}?)\s*\1",  # Repeated phrases
        "gibberish": r"[^a-zA-Z0-9\s.,\-()]",  # Non-standard characters
    }

    @staticmethod
    def validate_file_integrity(file_content: bytes, file_type: str) -> Tuple[bool, str]:
        """
        Validate file integrity and format.
        Returns: (is_valid, error_message)
        """
        logger.info(f"[validate_file_integrity] Checking {file_type} file integrity")
        
        try:
            if file_type == "pdf":
                return ResumeValidator._validate_pdf_integrity(file_content)
            elif file_type == "docx":
                return ResumeValidator._validate_docx_integrity(file_content)
            else:
                return False, "Invalid resume. Please upload a valid resume."
        except Exception as e:
            logger.error(f"[validate_file_integrity] Integrity check failed: {str(e)}")
            return False, "Invalid resume. Please upload a valid resume."

    @staticmethod
    def validate_resume_structure(file_content: bytes, file_type: str) -> Tuple[bool, str]:
        """
        Strong validation: Check if file actually contains resume content.
        This happens BEFORE parsing to quickly reject invalid files.
        Returns: (is_valid, error_message)
        """
        logger.info(f"[validate_resume_structure] Validating resume structure for {file_type}")
        
        try:
            if file_type == "pdf":
                text = ResumeValidator._extract_pdf_text(file_content)
            elif file_type == "docx":
                text = ResumeValidator._extract_docx_text(file_content)
            else:
                return False, "Invalid resume. Please upload a valid resume."
            
            # Check minimum content
            if len(text.strip()) < 100:
                return False, "Invalid resume. Please upload a valid resume."
            
            # Check for resume-like structure
            text_lower = text.lower()
            
            # Must have at least one of these sections
            required_sections = [
                ('experience', 'experience|employment|work|job'),
                ('education', 'education|degree|university|school|college'),
                ('skills', 'skill|expertise|technical|programming|languages'),
                ('contact', 'email|phone|contact|linkedin|github|website'),
            ]
            
            sections_found = []
            for section_name, section_pattern in required_sections:
                if re.search(section_pattern, text_lower):
                    sections_found.append(section_name)
            
            # Must have at least 2 resume sections
            if len(sections_found) < 2:
                return False, "Invalid resume. Please upload a valid resume."
            
            # Check for name/header (first line should have fewer than 15 words)
            first_line_words = text.split('\n')[0].split()
            if len(first_line_words) > 15:
                return False, "Invalid resume. Please upload a valid resume."
            
            # ✅ NEW: Check for obviously low-quality/fake resume indicators
            quality_issues = ResumeValidator._check_resume_quality(text)
            if quality_issues:
                return False, quality_issues
            
            logger.info(f"[validate_resume_structure] ✅ Resume structure valid. Sections found: {', '.join(sections_found)}")
            return True, ""
            
        except Exception as e:
            logger.error(f"[validate_resume_structure] Structure validation failed: {str(e)}")
            return False, "Invalid resume. Please upload a valid resume."

    @staticmethod
    def _check_resume_quality(text: str) -> str:
        """
        Check for obviously low-quality, fake, or obviously invalid resume indicators.
        Returns error message if found, empty string if quality is acceptable.
        """
        text_lower = text.lower()
        generic_error = "Invalid resume. Please upload a valid resume."
        
        # 1. Check for invalid email format
        email_pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
        emails = re.findall(email_pattern, text_lower)
        
        if emails:
            for email in emails:
                # Check for obvious fake emails
                if '@@' in email or '...' in email or '@@.' in email:
                    return generic_error
                if any(x in email for x in ['notanemail', 'fake', 'test', 'dummy', 'example@example']):
                    return generic_error
        
        # 2. Check for nonsense phone numbers
        phone_pattern = r'\b\d{7,}\b'
        phones = re.findall(phone_pattern, text)
        
        for phone in phones:
            # All same digits (0000000000, 1111111111, etc) is fake
            if len(set(phone)) == 1:
                return generic_error
            # Sequence pattern (12345678) is fake
            if ResumeValidator._is_sequential_number(phone):
                return generic_error
        
        # 3. Check for obviously unprofessional content
        unprofessional_keywords = [
            "typing fast sometimes",
            "knows computer on/off",
            "internet browsing",
            "sleeping",
            "scrolling",
            "youtube",
            "tiktok",
            "instagram",
            "works sometimes",
            "tried but not finished",
            "stuff",
            "things",
            "something",
        ]
        
        unprofessional_found = []
        for keyword in unprofessional_keywords:
            if keyword in text_lower:
                unprofessional_found.append(keyword)
        
        if len(unprofessional_found) >= 3:
            return generic_error
        elif len(unprofessional_found) > 0:
            # If we found some, check context - if it's in skills section, it's definitely bad
            if 'skills' in text_lower:
                skills_section = text_lower[text_lower.index('skills'):text_lower.index('skills') + 500]
                if any(kw in skills_section for kw in unprofessional_found):
                    return generic_error
        
        # 4. Check for degree name validity
        if 'degree' in text_lower or 'education' in text_lower:
            bad_degree_patterns = r'degree in (something|nothing|stuff|test|dummy)'
            if re.search(bad_degree_patterns, text_lower):
                return generic_error
        
        # 5. Check for URL/contact info in experience (sign of bad parsing or fake resume)
        exp_keywords = ['experience', 'employment', 'work', 'job']
        has_exp_section = any(kw in text_lower for kw in exp_keywords)
        
        if has_exp_section:
            # Experience should have company names and job titles, not URLs or obvious placeholders
            if 'somewhere pvt ltd' in text_lower and 'did something' in text_lower:
                return generic_error
        
        # 6. Check name validity
        lines = text.split('\n')
        first_line = lines[0].lower() if lines else ""
        
        # Name should not contain email or phone on same line
        if '@' in first_line or re.search(r'\d{3}', first_line):
            # Extract just the name part
            if 'name:' in first_line:
                name_part = first_line.split('name:')[1].strip()
                # Check if name has too many numbers/special chars
                if len(re.findall(r'\d+', name_part)) > 2:
                    return generic_error
        
        # Name like "ABC XYZ 123" with number is suspicious
        if re.search(r'\b[A-Z]{3,}\s+[A-Z]{3,}\s+\d{3,}\b', first_line):
            return generic_error
        
        logger.info("[_check_resume_quality] ✅ Resume quality checks passed")
        return ""  # All quality checks passed

    @staticmethod
    def _is_sequential_number(num_str: str) -> bool:
        """Check if phone number is a sequence (123456789, etc)."""
        try:
            # Check if digits are sequential
            digits = [int(d) for d in num_str]
            for i in range(len(digits) - 1):
                if abs(digits[i + 1] - digits[i]) > 1:
                    return False
            return True
        except:
            return False

    @staticmethod
    def _extract_pdf_text(file_content: bytes) -> str:
        """Extract all text from PDF."""
        try:
            pdf_reader = PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                try:
                    text += page.extract_text()
                except:
                    pass
            return text
        except:
            return ""

    @staticmethod
    def _extract_docx_text(file_content: bytes) -> str:
        """Extract all text from DOCX."""
        try:
            doc = Document(io.BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except:
            return ""

    @staticmethod
    def _validate_pdf_integrity(file_content: bytes) -> Tuple[bool, str]:
        """Validate PDF file integrity and content."""
        try:
            # Check PDF magic number
            if not file_content.startswith(b'%PDF'):
                logger.error("[_validate_pdf_integrity] Invalid PDF magic number")
                return False, "Invalid resume. Please upload a valid resume."
            
            # Try to read PDF
            pdf_reader = PdfReader(io.BytesIO(file_content))
            num_pages = len(pdf_reader.pages)
            
            if num_pages == 0:
                logger.error("[_validate_pdf_integrity] PDF has no pages")
                return False, "Invalid resume. Please upload a valid resume."
            
            # ✅ NEW: Extract and validate actual text content
            total_text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        total_text += page_text
                except Exception as e:
                    logger.warning(f"[_validate_pdf_integrity] Failed to extract text from page {page_num}: {str(e)}")
            
            # Check if PDF has actual text content (not just empty pages)
            text_length = len(total_text.strip())
            if text_length < 50:  # Minimum 50 characters of actual text
                logger.error(f"[_validate_pdf_integrity] PDF has no meaningful content (only {text_length} chars)")
                return False, "Invalid resume. Please upload a valid resume."
            
            logger.info(f"[_validate_pdf_integrity] ✅ Valid PDF with {num_pages} pages and {text_length} characters of content")
            return True, ""
            
        except Exception as e:
            logger.error(f"[_validate_pdf_integrity] PDF validation failed: {str(e)}")
            return False, "Invalid resume. Please upload a valid resume."

    @staticmethod
    def _validate_docx_integrity(file_content: bytes) -> Tuple[bool, str]:
        """Validate DOCX file integrity."""
        try:
            # Check DOCX magic number (ZIP format)
            if not file_content.startswith(b'PK'):
                logger.error("[_validate_docx_integrity] Invalid DOCX magic number")
                return False, "Invalid resume. Please upload a valid resume."
            
            # Try to read DOCX
            doc = Document(io.BytesIO(file_content))
            paragraphs = len(doc.paragraphs)
            
            if paragraphs == 0:
                logger.error("[_validate_docx_integrity] DOCX has no paragraphs")
                return False, "Invalid resume. Please upload a valid resume."
            
            logger.info(f"[_validate_docx_integrity] ✅ Valid DOCX with {paragraphs} paragraphs")
            return True, ""
            
        except Exception as e:
            logger.error(f"[_validate_docx_integrity] DOCX validation failed: {str(e)}")
            return False, "Invalid resume. Please upload a valid resume."

    @staticmethod
    def validate_content_quality(extracted_data: Dict) -> Tuple[bool, List[str]]:
        """
        Validate extracted resume content for quality and authenticity.
        Returns: (is_valid, list_of_warnings)
        """
        logger.info(f"[validate_content_quality] Analyzing extracted data")
        warnings = []
        
        # 1. Check if meaningful data was extracted
        fields_found = sum([
            bool(extracted_data.get('name')),
            bool(extracted_data.get('email')),
            bool(extracted_data.get('phone')),
            bool(extracted_data.get('skills')),
            bool(extracted_data.get('experiences')),
            bool(extracted_data.get('educations')),
        ])
        
        if fields_found < ResumeValidator.MIN_EXTRACTED_FIELDS:
            msg = f"Insufficient data extracted: only {fields_found} fields found (minimum {ResumeValidator.MIN_EXTRACTED_FIELDS} required)"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
            return False, warnings
        
        # 2. Check for contact information
        has_contact = bool(extracted_data.get('email')) or bool(extracted_data.get('phone'))
        if not has_contact:
            msg = "No contact information (email or phone) found in resume"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
        
        # 3. Check for professional content (skills or experience)
        skills = extracted_data.get('skills', [])
        experiences = extracted_data.get('experiences', [])
        educations = extracted_data.get('educations', [])
        
        has_professional_content = len(skills) > 0 or len(experiences) > 0 or len(educations) > 0
        if not has_professional_content:
            msg = "No professional content (skills, experience, or education) found"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
            return False, warnings
        
        # 4. Check for placeholder/template keywords
        full_text = ResumeValidator._extract_full_text(extracted_data).lower()
        fake_keywords_found = []
        for keyword in ResumeValidator.FAKE_RESUME_KEYWORDS:
            if keyword.lower() in full_text:
                fake_keywords_found.append(keyword)
        
        if fake_keywords_found:
            msg = f"Suspicious template keywords found: {', '.join(fake_keywords_found)}"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
            return False, warnings
        
        # 5. Validate name format
        name = extracted_data.get('name', '')
        if name:
            if not ResumeValidator._is_valid_name(name):
                msg = f"Invalid name format: {name}"
                logger.warning(f"[validate_content_quality] {msg}")
                warnings.append(msg)
        
        # 6. Validate email format
        email = extracted_data.get('email', '')
        if email and not ResumeValidator._is_valid_email(email):
            msg = f"Invalid email format: {email}"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
        
        # 7. Check email domain legitimacy
        if email and ResumeValidator._is_suspicious_email_domain(email):
            msg = f"Suspicious email domain: {email}"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
        
        # 8. Validate phone format
        phone = extracted_data.get('phone', '')
        if phone and not ResumeValidator._is_valid_phone(phone):
            msg = f"Invalid phone format: {phone}"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
        
        # 9. Check for minimum text length after extraction
        total_text_length = len(full_text)
        if total_text_length < ResumeValidator.MIN_TEXT_LENGTH:
            msg = f"Resume content too short: {total_text_length} characters (minimum {ResumeValidator.MIN_TEXT_LENGTH})"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
            return False, warnings
        
        # 10. Check for maximum text length (spam/garbage files)
        if total_text_length > ResumeValidator.MAX_TEXT_LENGTH:
            msg = f"Resume content too large: {total_text_length} characters (maximum {ResumeValidator.MAX_TEXT_LENGTH})"
            logger.warning(f"[validate_content_quality] {msg}")
            warnings.append(msg)
            return False, warnings
        
        logger.info(f"[validate_content_quality] ✅ Content quality passed ({fields_found} fields, {total_text_length} chars)")
        return True, warnings

    @staticmethod
    def detect_ai_generated_content(extracted_data: Dict) -> Tuple[bool, List[str]]:
        """
        Heuristic checks for AI-generated or synthetic resume content.
        Returns: (looks_genuine, list_of_suspicion_reasons)
        """
        logger.info(f"[detect_ai_generated_content] Analyzing for AI-generated patterns")
        suspicions = []
        
        # 1. Check for overly polished language (extremely common in AI)
        full_text = ResumeValidator._extract_full_text(extracted_data).lower()
        
        # Common AI buzzwords that appear in generated content
        ai_buzzwords = [
            "leveraging synergies",
            "paradigm shift",
            "cutting-edge technologies",
            "robust solutions",
            "seamless integration",
            "end-to-end",
            "best-in-class",
            "world-class",
        ]
        
        buzzword_count = sum(1 for buzzword in ai_buzzwords if buzzword in full_text)
        if buzzword_count >= 3:
            msg = f"Excessive use of buzzwords ({buzzword_count} found) - suggests AI-generated content"
            logger.warning(f"[detect_ai_generated_content] {msg}")
            suspicions.append(msg)
        
        # 2. Check for perfect grammar/structure (humans usually have minor errors)
        if ResumeValidator._has_suspiciously_perfect_grammar(full_text):
            msg = "Suspiciously perfect grammar and structure - unusual for human writing"
            logger.warning(f"[detect_ai_generated_content] {msg}")
            suspicions.append(msg)
        
        # 3. Check date consistency
        experiences = extracted_data.get('experiences', [])
        date_issues = ResumeValidator._check_date_consistency(experiences)
        if date_issues:
            suspicions.extend(date_issues)
        
        # 4. Check for generic descriptions
        generic_count = ResumeValidator._count_generic_descriptions(experiences)
        if generic_count > 0:
            msg = f"Found {generic_count} generic job descriptions - common in AI-generated content"
            logger.warning(f"[detect_ai_generated_content] {msg}")
            suspicions.append(msg)
        
        is_likely_genuine = len(suspicions) < 2  # Less than 2 suspicions = likely genuine
        logger.info(f"[detect_ai_generated_content] Genuineness confidence: {100 - (len(suspicions) * 25)}% (suspicions: {len(suspicions)})")
        
        return is_likely_genuine, suspicions

    # ============== Helper Methods ==============

    @staticmethod
    def _extract_full_text(extracted_data: Dict) -> str:
        """Extract all text from parsed resume data."""
        texts = [
            extracted_data.get('name', ''),
            extracted_data.get('email', ''),
            extracted_data.get('phone', ''),
            extracted_data.get('summary', ''),
        ]
        
        for skill in extracted_data.get('skills', []):
            texts.append(str(skill))
        
        for exp in extracted_data.get('experiences', []):
            texts.append(str(exp.get('title', '')))
            texts.append(str(exp.get('company', '')))
            texts.append(str(exp.get('description', '')))
        
        for edu in extracted_data.get('educations', []):
            texts.append(str(edu.get('degree', '')))
            texts.append(str(edu.get('institution', '')))
        
        return ' '.join(texts)

    @staticmethod
    def _is_valid_name(name: str) -> bool:
        """Validate name format."""
        if not name or len(name) < 2:
            return False
        if len(name) > 100:
            return False
        # Should have at least 2 words or contain letters
        if not re.search(r'[a-zA-Z]', name):
            return False
        return True

    @staticmethod
    def _is_valid_email(email: str) -> bool:
        """Validate email format."""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email))

    @staticmethod
    def _is_suspicious_email_domain(email: str) -> bool:
        """Check for suspicious email domains (temp emails, etc)."""
        suspicious_domains = [
            'tempmail', 'temp-mail', 'test.com', 'example.com',
            'demo.com', 'sample.com', 'placeholder.com', 'fake.com'
        ]
        domain = email.split('@')[1].lower() if '@' in email else ''
        return any(suspicious in domain for suspicious in suspicious_domains)

    @staticmethod
    def _is_valid_phone(phone: str) -> bool:
        """Validate phone format."""
        # Remove common formatting
        cleaned = re.sub(r'[\s\-().]', '', phone)
        # Should have 7-15 digits
        return bool(re.match(r'^\+?1?\d{7,15}$', cleaned))

    @staticmethod
    def _has_suspiciously_perfect_grammar(text: str) -> bool:
        """Check for suspiciously perfect grammar (AI indicator)."""
        # Simple heuristic: humans usually have at least some capitalization issues
        sentences = text.split('.')
        if len(sentences) < 3:
            return False
        
        # Count sentences with proper capitalization
        proper_sentences = sum(1 for s in sentences if s and s[0].isupper())
        ratio = proper_sentences / len(sentences) if sentences else 0
        
        # More than 90% perfect capitalization is suspicious
        return ratio > 0.9

    @staticmethod
    def _check_date_consistency(experiences: List[Dict]) -> List[str]:
        """Check for date inconsistencies in experiences."""
        issues = []
        
        for i, exp in enumerate(experiences):
            start_date = exp.get('start_date', '')
            end_date = exp.get('end_date', '')
            
            if start_date and end_date:
                try:
                    # Extract years
                    start_year = int(re.search(r'\d{4}', start_date).group())
                    end_year = int(re.search(r'\d{4}', end_date).group())
                    
                    if start_year > end_year:
                        issues.append(f"Experience {i+1} has end date before start date")
                    
                    duration = end_year - start_year
                    if duration > 50:
                        issues.append(f"Experience {i+1} duration unusually long ({duration} years)")
                except:
                    pass
        
        return issues

    @staticmethod
    def _count_generic_descriptions(experiences: List[Dict]) -> int:
        """Count generic/template job descriptions."""
        generic_phrases = [
            "responsible for",
            "contributed to",
            "involved in",
            "worked on",
            "part of the team",
            "helped with",
        ]
        
        count = 0
        for exp in experiences:
            description = (exp.get('description', '') or '').lower()
            if description and sum(1 for phrase in generic_phrases if phrase in description) >= 2:
                count += 1
        
        return count
